import json
import os
from datetime import datetime
from django.conf import settings
from django.core.files.storage import default_storage
from django.forms import model_to_dict
from django.shortcuts import get_object_or_404
from html2docx import html2docx
from .models import *

from django.core.paginator import Paginator
from django.db.models import Q, Sum, CharField, Count
from django.middleware.csrf import get_token
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.db.models import Case, When, Value
import pandas as pd
from openai import OpenAI
import docx
from docx import Document

from .toolsword import word_to_pdf, save_content_as_word, pdf_to_word


def Fail(data=None):
    if data is not None:
        return JsonResponse({'code': 400, 'msg': '失败', 'data': data})
    else:
        return JsonResponse({'code': 400, 'msg': '失败'})


def Suc(data=None, total=0):
    if data is not None:
        return JsonResponse({'code': 200, 'msg': '成功', 'data': data, 'total': total})
    else:
        return JsonResponse({'code': 200, 'msg': '成功', 'data': None, 'total': total})


# 注册
@csrf_exempt
def add_user(request):
    if request.method == 'POST':
        # 从请求体中获取原始数据
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return Fail()
        print(data)
        # 获取用户数据
        username = data.get('username')
        name = data.get('name', None)
        password = data.get('password')
        role_id = data.get('role', 1)
        sex = int(data.get('sex', 1))
        birth_date = data.get('birth_date', "")
        print(birth_date)
        if birth_date:
            try:
                # 尝试匹配包含时间部分的格式
                date_time_obj = datetime.strptime(birth_date, "%Y-%m-%dT%H:%M:%S.%fZ")
            except ValueError:
                # 如果不含时间部分，则使用简单日期格式
                date_time_obj = datetime.strptime(birth_date, "%Y-%m-%d")

        phone = data.get('phone', 1)

        # 创建用户对象并保存到数据库
        try:
            user = User.objects.create(
                username=username,
                name=name,
                password=password,
                role=role_id,
                sex=sex,
                birth_date=date_time_obj,
                phone=phone

            )
            # class_query = Class.objects.filter(id=classId)
            # class_obj = class_query.first()
            # class_= ClassStudent.objects.create(
            #
            # class_id=class_obj,
            # student =user
            #
            # )

            return Suc()
        except Exception as e:
            print(e)
            return Fail()
    else:
        return Fail()


@csrf_exempt
def user_login(request):
    if request.method == 'POST':
        # 从请求体中获取原始数据
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return Fail()

        # 获取登录表单中的数据
        username = data.get('username')
        password = data.get('password')

        # 查询数据库，检查是否存在具有特定用户名和密码的用户
        user_query = User.objects.filter(username=username, password=password)

        # 检查查询结果是否存在
        if user_query.exists():

            user_obj = user_query.first()
            # 构建要返回的用户信息
            user_info = {
                'id': user_obj.id,
                'username': user_obj.username,
                'name': user_obj.name,
                'password': user_obj.password,
                'role': user_obj.role,
                'sex': user_obj.sex,
                'phone': user_obj.phone,
                'birth_date': user_obj.birth_date,

            }
            res = {'user': user_info, 'token': get_token(request)}  # X-CSRFTOKEN
            # 用户存在，返回登录成功的 JSON 响应
            return Suc(res)
        else:
            # 用户不存在，返回登录失败的 JSON 响应
            return Fail()

    else:
        # 如果请求方法不是 POST，则返回错误信息
        return Fail()


# 账号唯一检查
@csrf_exempt
def user_check(request):
    username = request.GET.get('username')
    # 查询数据库，检查是否存在具有特定用户名和密码的用户
    user_query = User.objects.filter(username=username)

    # 检查查询结果是否存在
    if user_query.exists():
        user_obj = user_query.first()
        # 构建要返回的用户信息
        user_info = {
            'username': user_obj.username,
            'name': user_obj.name,
            'password': user_obj.password,
            'role': user_obj.role,
            'sex': user_obj.sex,

        }
        return Fail(user_info)
    else:
        # 用户不存在，返回登录失败的 JSON 响应
        return Suc()


@csrf_exempt
def delete_user(request):
    id = int(request.GET.get('id'))
    print(id)
    try:
        # 删除用户

        user = User.objects.get(id=id)
        print(user)
        user.delete()

        return Suc()
    except Exception as e:
        print(e)
        return Fail()


@csrf_exempt
def generate_file_url(request, file_path):
    """
    根据静态文件的路径，生成并返回文件的 URL。
    :param file_path: 文件相对于 STATIC_URL 的路径
    :return: JsonResponse 返回文件的 URL
    """
    # 拼接静态文件的 URL
    file_url = os.path.join(settings.STATIC_URL, file_path)

    # 生成并返回文件的完整 URL
    return JsonResponse({"file_url": file_url})


# 用户改
@csrf_exempt
def update_user(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            id = data.get('id')

            user = User.objects.get(id=id)
        except (json.JSONDecodeError, User.DoesNotExist):
            return Fail()

        # 更新用户数据
        user.name = data.get('name', user.name)
        user.username = data.get('username', user.username)
        user.password = data.get('password', user.password)
        user.role = int(data.get('role', user.role))  # 转换为整数
        user.sex = int(data.get('sex', user.sex))  # 转换为整数

        # 处理出生日期
        birth_date_str = data.get('birth_date', None)
        print(birth_date_str)
        if birth_date_str:
            try:
                # 尝试匹配包含时间部分的格式
                date_time_obj = datetime.strptime(birth_date_str, "%Y-%m-%dT%H:%M:%S.%fZ")
            except ValueError:
                # 如果不含时间部分，则使用简单日期格式
                date_time_obj = datetime.strptime(birth_date_str, "%Y-%m-%d")
            user.birth_date = date_time_obj

        user.phone = data.get('phone', user.phone)

        # 保存更新后的用户信息
        user.save()
        return Suc()


# list
@csrf_exempt
def user_list(request):
    if request.method == 'POST':
        # 从请求体中获取原始数据
        try:
            data = json.loads(request.body)
            name = data.get('name', None)
            page_size = data.get('pageSize', 10)  # 默认每页显示10条数据
            page_num = data.get('pageNum', 1)  # 默认显示第一页
        except json.JSONDecodeError:
            return Fail()

    if name is not None:
        users = User.objects.filter(name__icontains=name)
    else:
        users = User.objects.all()

    paginator = Paginator(users, page_size)
    page_obj = paginator.get_page(page_num)

    # 将分页后的数据转换为字典列表
    users_data = list(page_obj.object_list.values())

    return Suc(users_data, paginator.count)


# 项目

@csrf_exempt
def add_project(request):
    if request.method == 'POST':
        # 从请求体中获取原始数据
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return Fail()
        print(data)
        # 获取项目数据
        name = data.get('name')
        folder_path = data.get('folder_path', None)
        directory_image_url = data.get('directory_image_url')

        user_data = data.get('username')  # 获取完整的用户字典

        # 确保获取到用户字典
        if user_data:

            try:
                user = User.objects.get(username=user_data)  # 根据 ID 获取用户对象
            except User.DoesNotExist:
                return Fail()
        else:
            return Fail()

        # 创建项目对象并保存到数据库
        try:

            project = Project.objects.create(
                name=name,
                folder_path=folder_path,
                directory_image_url=directory_image_url,

                user=user,  # 将用户对象作为外键传入
            )
            print(123123)
            ProjectUserAssociation.objects.create(project=project, user=user)
            return Suc()
        except Exception as e:
            print(e)
            return Fail()
    else:
        return Fail()


@csrf_exempt
def delete_project(request):
    id = request.GET.get('id')
    try:
        # 删除用户

        project = Project.objects.get(id=id)

        project.delete()

        return Suc()
    except Exception as e:

        return Fail()


@csrf_exempt
def update_project(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            project_id = data.get('id')  # 获取项目 ID

            # 获取项目对象
            project = Project.objects.get(id=project_id)
        except (json.JSONDecodeError, Project.DoesNotExist):
            return JsonResponse({"status": "fail", "message": "项目不存在或请求数据无效"}, status=400)

        # 更新项目数据
        project.name = data.get('name', project.name)
        project.directory_image_url = data.get('directory_image_url', project.directory_image_url)

        user_data = data.get('user_name')
        if user_data:
            try:
                # 获取用户对象
                user = User.objects.get(username=user_data)

                # 检查 ProjectUserAssociation 是否存在
                association = ProjectUserAssociation.objects.filter(user=user, project=project).first()
                if association:
                    # 如果关联存在，更新项目管理员
                    project.user = user
                else:
                    return JsonResponse({"status": "fail", "message": "该用户不是该项目的合法伙伴"}, status=400)

            except User.DoesNotExist:
                return JsonResponse({"status": "fail", "message": "用户不存在"}, status=400)
        else:
            return JsonResponse({"status": "fail", "message": "用户名称不能为空"}, status=400)

        # 保存更新后的项目信息
        project.save()
        return Suc()


@csrf_exempt
def project_list(request):
    if request.method == 'POST':
        # 从请求体中获取原始数据
        try:
            data = json.loads(request.body)
            print(data)
            name = data.get('name', None)
            user = data.get('user', None)
            print(user)
            page_size = data.get('pageSize', 10)  # 默认每页显示10条数据
            page_num = data.get('pageNum', 1)  # 默认显示第一页
        except json.JSONDecodeError:
            return Fail()  # 解析失败时返回错误响应

    # 判断是否提供用户信息
    if user is not None:
        print(11111111)
        # 检查用户角色是否有效
        if 'role' in user and user["role"] == '0':
            print(222222222)
            # 管理员查询
            if name is not None:
                projects = Project.objects.filter(name__icontains=name).select_related('user')
            else:
                projects = Project.objects.all().select_related('user')
        else:
            # 普通用户查询
            if 'id' in user:
                try:
                    user_instance = User.objects.get(id=user["id"])
                except User.DoesNotExist:
                    return JsonResponse({'error': 'User not found'}, status=404)

                # 获取与该用户相关联的所有项目
                project_associations = ProjectUserAssociation.objects.filter(user=user_instance)
                print(project_associations)
                if project_associations.exists():
                    print(111111)
                    if name is not None and name != "":
                        projects = Project.objects.filter(
                            id__in=project_associations.values_list('project_id', flat=True), name__icontains=name
                        ).select_related('user')
                    else:
                        print(12312)
                        projects = Project.objects.filter(
                            id__in=project_associations.values_list('project_id', flat=True)
                        ).select_related('user')
                        print(projects)
                else:
                    projects = []  # 没有关联的项目，返回空列表

            else:
                return Fail()  # 如果没有用户ID，返回失败

    else:
        return Fail()  # 如果没有用户信息，返回失败

    # 如果没有项目，直接返回空列表
    if not projects:
        return JsonResponse({'data': [], 'total': 0})  # 无内容时返回空数组
    print(projects)
    # 分页处理
    paginator = Paginator(projects, page_size)
    page_obj = paginator.get_page(page_num)

    # 获取分页后的数据
    projects_data = []
    print(len(page_obj.object_list))
    for project in page_obj.object_list:
        # 提取项目字段并获取相关联的用户数据
        project_data = {
            'id': project.id,
            'name': project.name,
            'directory_image_url': project.directory_image_url,
            'created_at': project.created_at,
            'folder_path': project.folder_path,
            'updated_at': project.updated_at,

            'user_id': project.user.id if project.user else None,  # 获取 user_id
            'username': project.user.username if project.user else None,  # 获取用户名
            'user_name': project.user.name if project.user else None,  # 获取用户姓名
            # 其他需要的字段
        }
        print(project_data)
        projects_data.append(project_data)

    # 返回成功响应，传递分页数据和总计数
    return Suc(projects_data, paginator.count)


@csrf_exempt
def project_user_delete(request):
    if request.method == 'POST':
        try:
            # 解析传入的 JSON 数据
            data = json.loads(request.body)
            project_id = data.get('project_id')

            user_username = data.get('username')

            # 校验参数是否完整
            if not project_id or not user_username:
                return Fail()

            # 获取用户、用户主信息和项目
            user = get_object_or_404(User, username=user_username)

            project = get_object_or_404(Project, id=project_id)

            # 尝试删除用户与项目的关联
            association = ProjectUserAssociation.objects.filter(project=project, user=user).first()
            print(1212)
            if association:
                print(association)
                # 删除关联
                association.delete()
                return Suc()
            else:
                return Fail()

        except json.JSONDecodeError:
            return Fail()
        except Exception as e:
            return Fail()

    # 如果请求方法不是 POST，返回方法错误
    return Fail()


@csrf_exempt
def project_user_add(request):
    if request.method == 'POST':
        try:
            # 解析传入的 JSON 数据
            data = json.loads(request.body)
            project_id = data.get('project_id')

            user_username = data.get('username')

            # 校验参数是否完整
            if not project_id or not user_username:
                return JsonResponse({'error': '缺少必要的参数'}, status=400)

            # 获取用户、用户主信息和项目
            user = get_object_or_404(User, username=user_username)
            print(user.id)
            project = get_object_or_404(Project, id=project_id)
            print(project.id)
            # 创建用户与项目的关联，如果已存在则不会创建
            project_user_association, created = ProjectUserAssociation.objects.get_or_create(project=project, user=user)

            # 如果成功创建关联，返回成功信息；如果关联已经存在，则返回提示信息
            if created:
                print(1111111111)
                return Suc()
            else:

                return Fail()

        except json.JSONDecodeError:
            return Fail()
        except Exception as e:
            return Fail()

    # 如果请求方法不是 POST，返回方法错误
    return Fail()


@csrf_exempt
def project_user_list(request):
    if request.method == 'POST':
        try:
            print(1111111111)
            # 解析传入的 JSON 数据
            data = json.loads(request.body)
            print(data)

            project_id = data.get('project_id')

            # 校验参数是否完整
            if not project_id:
                return Fail()

            # 获取项目
            project = get_object_or_404(Project, id=project_id)

            # 获取项目相关的用户关联列表
            project_user_associations = ProjectUserAssociation.objects.filter(project=project)

            # 如果有用户关联，返回用户信息
            if project_user_associations:
                user_data = [
                    {
                        'id': assoc.user.id,
                        'username': assoc.user.username,
                        'phone': assoc.user.phone,  # 假设 User 模型有 email 字段
                        'name': assoc.user.name,  # 假设 User 模型有 name 字段
                        # 可以添加更多的用户字段
                    }
                    for assoc in project_user_associations
                ]

                return Suc(data=user_data)
            else:
                return Fail()

        except json.JSONDecodeError:
            return Fail()
        except Exception as e:
            return Fail()

    # 如果请求方法不是 POST，返回方法错误
    return Fail()


@csrf_exempt
def filelist(request):
    print(121212)
    if request.method == 'POST':
        try:
            # 从请求体中获取数据
            data = json.loads(request.body)
            print(data)
            project_id = data.get('project_id', None)

            # 如果 project_id 存在，获取与之关联的所有文件
            if project_id is not None:
                # 查询与 project_id 关联的所有 file 对象，并关联查询 project 和 project 的 user 信息
                files = File.objects.filter(
                    id__in=FileProjectAssociation.objects.filter(project_id=project_id)
                    .values_list('file', flat=True)
                ).select_related('user')  # 通过 select_related 来优化 user 查询

                # 获取与 project 相关的 user 信息
                project_user = Project.objects.filter(id=project_id).values('user__id', 'user__name').first()

                # 遍历 files，检查是否有与 file 关联的 Fileing 记录
                files_data = []
                for file in files:
                    # 查询 Fileing 表是否存在与该 file 的关联记录
                    fileing_record = Fileing.objects.filter(file=file).first()  # 使用 first() 获取第一个匹配的记录，避免多次查询

                    # 如果关联记录存在，获取 fileing 的 user_id，否则返回 None
                    if fileing_record:
                        fileing_exists = 1
                        fileing_user_id = fileing_record.user.id  # 获取 user 的 id
                    else:
                        fileing_exists = 0
                        fileing_user_id = None

                    # 将文件信息和 fileing_status 加入到返回的数据中，同时添加 project_user
                    files_data.append({
                        'id': file.id,
                        'name': file.name,
                        'file_url': file.file_url,
                        'file_path': file.file_path,
                        'created_at': file.created_at,
                        'updated_at': file.updated_at,
                        'user__name': file.user.name,
                        'fileing_status': fileing_exists,  # 新增的 fileing 状态
                        'fileing_user_id': fileing_user_id,  # 返回 fileing 关联的 user id
                        'project_user_id': project_user['user__id'],  # 项目用户 ID
                        'project_user_name': project_user['user__name'],  # 项目用户姓名
                    })
                # 直接返回合并后的数据
                return Suc(files_data)

            else:
                return Fail()

        except json.JSONDecodeError:
            return Fail()


# 正在编辑文件
@csrf_exempt
def fileing(request):
    if request.method == 'POST':
        try:
            # 从请求体中获取数据
            data = json.loads(request.body)

            file_id = data.get('file_id', None)

            # 如果 project_id 存在，获取与之关联的所有文件
            if file_id is not None:

                file = File.objects.get(id=file_id)
                # 查询与 project_id 关联的所有 file 对象
                if (Fileing.objects.filter(file=file).exists()):

                    return Suc(1)

                else:
                    return Fail()
            else:
                return Fail()
        except json.JSONDecodeError:
            return Fail()


@csrf_exempt
def fileing_add(request):
    if request.method == 'POST':
        try:
            print(json.loads(request.body))
            # 从请求体中获取数据
            data = json.loads(request.body)
            user_id = data.get('user_id', None)
            file_id = data.get('file_id', None)

            # 如果 user_id 和 file_id 都存在
            if user_id is not None and file_id is not None:
                try:
                    # 获取用户和文件对象
                    user = User.objects.get(id=user_id)
                    file = File.objects.get(id=file_id)

                    # 检查是否已经存在相同的关联
                    if Fileing.objects.filter(user=user, file=file).exists():
                        return JsonResponse({'status': 'fail', 'message': 'User is already editing this file.'})

                    # 创建新的关联记录
                    fileing = Fileing.objects.create(user=user, file=file)
                    print(12312)
                    return Suc()

                except User.DoesNotExist:
                    return Fail()
                except File.DoesNotExist:
                    return Fail()

            else:
                return Fail()

        except json.JSONDecodeError:
            return Fail()


@csrf_exempt
def file_del(request):
    if request.method == 'GET':
        try:
            # 从请求参数中获取数据

            file_id = request.GET.get('id')

            # 如果  file_id 存在
            if file_id is not None:
                try:

                    # 查找是否存在该用户和文件的关联
                    fileing = File.objects.filter(id=file_id).first()

                    if fileing:
                        os.remove(fileing.file_path)

                        # 删除该记录
                        fileing.delete()
                        return Suc()
                    else:
                        return Fail()

                except User.DoesNotExist:
                    return Fail()
                except File.DoesNotExist:
                    return Fail()

            else:
                return Fail()

        except json.JSONDecodeError:
            return Fail()


@csrf_exempt
def fileing_del(request):
    if request.method == 'GET':
        try:
            # 从请求参数中获取数据

            file_id = request.GET.get('id')

            # 如果  file_id 存在
            if file_id is not None:
                try:

                    file = File.objects.get(id=file_id)

                    # 查找是否存在该用户和文件的关联
                    fileing = Fileing.objects.filter(file=file).first()

                    if fileing:

                        # 删除该记录
                        fileing.delete()
                        return Suc()
                    else:
                        return Fail()

                except User.DoesNotExist:
                    return Fail()
                except File.DoesNotExist:
                    return Fail()

            else:
                return Fail()

        except json.JSONDecodeError:
            return Fail()


@csrf_exempt
def fileupload(request):
    if request.method == 'POST':
        try:
            print(request.POST)
            # 获取 user_id 和 project_id 参数
            user_id = request.POST.get('user_id')
            project_id = request.POST.get('project_id')

            # 获取上传的文件
            uploaded_file = request.FILES.get('file')

            # 验证是否存在 user_id, project_id 和文件
            if user_id and project_id and uploaded_file:
                try:
                    # 获取 User 和 Project 对象
                    user = User.objects.get(id=user_id)
                    project = Project.objects.get(id=project_id)

                    # 设置文件保存路径
                    static_dir = os.path.join(settings.BASE_DIR, 'static', 'uploads')  # static/uploads 文件夹
                    if not os.path.exists(static_dir):
                        os.makedirs(static_dir)  # 如果目录不存在，创建目录

                    # 检查文件是否已存在，若存在则修改文件名

                    file_path = os.path.join(static_dir, uploaded_file.name)

                    # 检查文件是否已经存在，如果存在则抛出异常
                    if os.path.exists(file_path):
                        return Fail("名称已存在")

                    # 保存文件到指定路径
                    with open(file_path, 'wb') as f:
                        for chunk in uploaded_file.chunks():
                            f.write(chunk)

                    # 获取文件的 URL
                    file_url = os.path.join('static', 'uploads', uploaded_file.name)

                    # 保存文件信息到 File 表
                    file = File.objects.create(
                        name=uploaded_file.name,
                        file_url="http://127.0.0.1:8000/" + file_url,
                        file_path=file_path,
                        user=user
                    )

                    # 保存文件与项目的关联信息
                    FileProjectAssociation.objects.create(
                        file=file,
                        project=project
                    )

                    return Suc()  # 返回成功响应

                except User.DoesNotExist:
                    return JsonResponse({'status': 'fail', 'message': 'User not found.'}, status=404)
                except Project.DoesNotExist:
                    return JsonResponse({'status': 'fail', 'message': 'Project not found.'}, status=404)

            else:
                return JsonResponse({'status': 'fail', 'message': 'Missing user_id, project_id, or file.'}, status=400)

        except Exception as e:
            # 捕获所有异常并返回错误信息
            return JsonResponse({'status': 'fail', 'message': str(e)}, status=500)


@csrf_exempt
def jpgupload(request):
    if request.method == 'POST':
        try:
            print(request.body)

            # 获取上传的文件
            uploaded_file = request.FILES.get('file')

            # 验证是否存在用户、项目和文件
            if uploaded_file:
                try:
                    # 获取用户和项目对象

                    # 设置文件保存路径，保存到 static/uploads 文件夹中
                    static_dir = os.path.join(settings.BASE_DIR, 'static', 'uploads')  # static/uploads 文件夹
                    if not os.path.exists(static_dir):
                        os.makedirs(static_dir)  # 如果目录不存在，创建目录

                    file_path = os.path.join(static_dir, uploaded_file.name)

                    # 检查文件是否已经存在，如果存在则抛出异常
                    if os.path.exists(file_path):
                        return Fail("名称已存在")

                    # 使用 default_storage 保存文件
                    with open(file_path, 'wb') as f:
                        for chunk in uploaded_file.chunks():
                            f.write(chunk)

                    # 获取文件的 URL
                    file_url = os.path.join('static', 'uploads', uploaded_file.name)

                    # 返回文件的 URL 和路径
                    return JsonResponse({
                        'status': 'success',
                        'message': 'File uploaded successfully.',
                        'file_url': "http://127.0.0.1:8000/" + file_url,
                        'file_path': file_path
                    })

                except User.DoesNotExist:
                    return Fail({'status': 'fail', 'message': 'User not found.'}, status=404)
                except Project.DoesNotExist:
                    return Fail({'status': 'fail', 'message': 'Project not found.'}, status=404)

            else:
                return Fail({'status': 'fail', 'message': 'Missing user_id, project_id, or file.'}, status=400)

        except json.JSONDecodeError:
            return Fail({'status': 'fail', 'message': 'Invalid JSON data.'}, status=400)

    return Fail({'status': 'fail', 'message': 'Invalid request method.'}, status=405)


from django.http import JsonResponse
from .models import File
import json
import mammoth


# 文件显示
@csrf_exempt
def fileinlist(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            file_id = data.get('file_id', None)
            type = data.get('type', None)

            if file_id is not None:
                file = File.objects.filter(id=file_id).first()
                if file:
                    file_path = file.file_path  # 假设文件存储路径在 file.file.path 中
                    html_content = ''
                    if type == "word":
                        # 使用 mammoth 将 .docx 转换为 HTML
                        with open(file_path, 'rb') as docx_file:
                            result = mammoth.convert_to_html(docx_file)
                            html_content = result.value  # 获取转换后的 HTML 内容

                    # 将文件信息和 HTML 内容一起返回
                    file_data = {
                        'id': file.id,
                        'name': file.name,
                        'file_url': file.file_url,
                        'file_path': file_path,  # 可以根据需要返回更多的文件信息
                        'html_content': html_content,
                    }

                    return JsonResponse({'status': 'success', 'data': file_data})
                else:
                    return JsonResponse({'status': 'fail', 'message': 'File not found'}, status=404)
            else:
                return JsonResponse({'status': 'fail', 'message': 'No file_id provided'}, status=400)

        except json.JSONDecodeError:
            return JsonResponse({'status': 'fail', 'message': 'Invalid JSON'}, status=400)


# word文档
@csrf_exempt
def save_content(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            content = data.get('content', None)
            file_id = data.get('file_id', None)

            if file_id and content:
                # 获取文件对象
                file = File.objects.filter(id=file_id).first()

                # 设置文件保存路径
                static_dir = os.path.join(settings.BASE_DIR, 'static', 'uploads')  # static/uploads 文件夹
                if not os.path.exists(static_dir):
                    os.makedirs(static_dir)  # 如果目录不存在，创建目录

                # 使用html2docx库将HTML内容转换为Word格式
                docx_file_path = os.path.join(static_dir, file.name)
                # 使用html2docx将HTML内容转换为BytesIO对象
                buf = html2docx(content, title="My Document")

                # 将字节流写入文件
                with open(docx_file_path, "wb") as docx_file:
                    docx_file.write(buf.getvalue())

                return JsonResponse({'message': 'File saved as Word document.', 'file_path': docx_file_path})

        except Exception as e:
            print(e)
            return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
def save_as_pdf(request):
    if request.method == 'POST':
        try:
            # 从请求体中获取数据
            data = json.loads(request.body)
            content = data.get('content', None)  # 富文本内容
            file_id = data.get('file_id', None)
            user_id = data.get('user_id', None)

            if file_id and content:

                file = File.objects.filter(id=file_id).first()
                project = FileProjectAssociation.objects.filter(file=file)
                print(123112)
                # 设置文件保存路径
                static_dir = os.path.join(settings.BASE_DIR, 'static', 'uploads')  # static/uploads 文件夹
                if not os.path.exists(static_dir):
                    os.makedirs(static_dir)  # 如果目录不存在，创建目录

                base, ext = os.path.splitext(file.name)

                # 设置 Word 和 PDF 文件路径
                word_file_path = os.path.join(static_dir, '123122ew313' + ".docx")
                pdf_file_path = os.path.join(static_dir, base + ".pdf")

                # 保存 Word 文件
                save_content_as_word(content, word_file_path)

                # # 删除 Word 文件
                # if os.path.exists(word_file_path):
                #     os.remove(word_file_path)

                # 获取文件的 URL
                file_url = os.path.join('static', 'uploads', base + ".pdf")
                print(123123)
                # 如果 PDF 文件存在，更新或覆盖现有记录
                if os.path.exists(pdf_file_path):
                    file = File.objects.get(file_path=pdf_file_path)
                    file.file_url = "http://127.0.0.1:8000/" + file_url
                    file.file_path = pdf_file_path
                    file.name = base + ".pdf"
                    file.save()

                else:
                    # 创建新的文件记录
                    file = File.objects.create(
                        name=base + ".pdf",
                        file_url="http://127.0.0.1:8000/" + file_url,
                        file_path=pdf_file_path,
                        user=User.objects.filter(id=user_id).first()
                    )
                    print(file)
                    # 保存文件与项目的关联信息
                    FileProjectAssociation.objects.create(
                        file=file,
                        project=project.first().project  # 取第一个关联项目
                    )
                    # 转换 Word 文件到 PDF
                word_to_pdf(word_file_path, pdf_file_path)

                return Suc()

            else:
                return JsonResponse({'status': 'fail', 'message': 'Missing user_id, file_id or content.'}, status=400)

        except Exception as e:
            print(e)
            # 捕获所有异常并返回错误信息
            return JsonResponse({'status': 'fail', 'message': str(e)}, status=500)


# pdf-》word
@csrf_exempt
def save_as_word(request):
    if request.method == 'POST':
        try:
            # 从请求体中获取数据
            data = json.loads(request.body)

            file_id = data.get('file_id', None)
            user_id = data.get('user_id', None)

            if file_id and user_id:
                file = File.objects.filter(id=file_id).first()
                project = FileProjectAssociation.objects.filter(file=file)

                # 设置文件保存路径
                static_dir = os.path.join(settings.BASE_DIR, 'static', 'uploads')  # static/uploads 文件夹
                if not os.path.exists(static_dir):
                    os.makedirs(static_dir)  # 如果目录不存在，创建目录

                base, ext = os.path.splitext(file.name)

                # 将扩展名改为 .pdf
                new_file_path = base + ".docx"

                file_path = os.path.join(static_dir, new_file_path)

                # 获取文件的 URL
                file_url = os.path.join('static', 'uploads', new_file_path)

                if os.path.exists(file_path):
                    pass

                else:
                    # 创建新的文件记录
                    file = File.objects.create(
                        name=base + ".docx",
                        file_url="http://127.0.0.1:8000/" + file_url,
                        file_path=file_path,
                        user=User.objects.filter(id=user_id).first()
                    )
                    print()
                    # 保存文件与项目的关联信息
                    FileProjectAssociation.objects.create(
                        file=file,
                        project=project.first().project  # 取第一个关联项目
                    )
                pdf_to_word(f"{file.file_path}", file_path)
                return JsonResponse({'status': 'success', 'message': 'Content saved as PDF successfully'})

            else:
                return JsonResponse({'status': 'fail', 'message': 'Missing user_id, file_id or content.'}, status=400)

        except Exception as e:
            # 捕获所有异常并返回错误信息
            return JsonResponse({'status': 'fail', 'message': str(e)}, status=500)


@csrf_exempt
def execl_get(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            file_id = data.get('file_id')
            print()

            if file_id:
                # 获取文件路径
                file = File.objects.get(id=file_id)
                file_path = file.file_path

                # 读取 Excel 文件内容
                df = pd.read_excel(file_path)
                df = df.fillna('')  # 替换 NaN 为空字符串
                excel_data = df.to_dict(orient='records')
                # 返回数据
                return Suc({
                    'status': 'success',
                    'data': excel_data,
                    'file': file.file_url
                })
            else:
                return JsonResponse({'status': 'error', 'message': 'file_id not provided'}, status=400)
        except Exception as e:
            print(e)
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
    else:
        return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=405)


@csrf_exempt
def execl_set(request):
    if request.method == 'POST':
        try:
            # 接收前端数据
            data = json.loads(request.body)
            updated_data = data.get('data', [])

            if not updated_data:
                return JsonResponse({'status': 'error', 'message': 'No data provided'}, status=400)

            # 获取对应的 Excel 文件
            file_id = data.get('file_id')  # 确保前端发送了 file_id
            file = File.objects.get(id=file_id)
            file_path = file.file_path

            # 将数据写回 Excel 文件
            df = pd.DataFrame(updated_data)
            df.to_excel(file_path, index=False)  # 保存到 Excel

            return Suc({'status': 'success', 'message': 'Data updated successfully'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
    else:
        return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=405)


@csrf_exempt
def generate_exam(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            project_id = data.get('project_id')
            user_id = data.get('user_id')

            if not project_id or not user_id:
                return Fail("缺少必要参数")

            # 获取用户和项目
            try:
                user = User.objects.get(id=user_id)
                project = Project.objects.get(id=project_id)
            except (User.DoesNotExist, Project.DoesNotExist):
                return Fail("用户或项目不存在")

            # 获取项目中的所有文件
            file_associations = FileProjectAssociation.objects.filter(project_id=project_id)
            files = [assoc.file for assoc in file_associations]

            if not files:
                return Fail("项目中没有文件")

            # 提取所有文件的内容
            doc_content = ""
            for file in files:
                file_path = file.file_path
                file_name = file.name.lower()

                try:
                    # 根据文件类型提取内容
                    if file_name.endswith('.docx'):
                        doc = Document(file_path)
                        content = "\n".join([para.text for para in doc.paragraphs if para.text])
                        doc_content += f"文件名: {file.name}\n内容:\n{content}\n\n"
                    elif file_name.endswith('.pdf'):
                        # 这里需要添加PDF内容提取的代码
                        # 可以使用PyPDF2或pdfplumber等库
                        doc_content += f"文件名: {file.name}\n内容: [PDF内容]\n\n"
                    elif file_name.endswith('.xlsx'):
                        # 这里需要添加Excel内容提取的代码
                        df = pd.read_excel(file_path)
                        content = df.to_string(index=False)
                        doc_content += f"文件名: {file.name}\n内容:\n{content}\n\n"
                except Exception as e:
                    print(f"提取文件 {file.name} 内容时出错: {str(e)}")
                    continue

            if not doc_content:
                return Fail("无法提取文件内容")

            # 调用AI生成试题
            system_prompt = """
            你是一个专业的教育内容分析和试题生成专家。请根据提供的文档内容，生成高质量的试题。
            你的输出应该是一个完整的试卷，包含多种题型（如选择题、填空题、简答题等）。
            请确保试题与文档内容紧密相关，难度适中，并且格式规范。
            """

            # 初始化AI客户端
            client = OpenAI(
                api_key="sk-f7728320262b4caf951dd8087e1923d6",
                base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
            )

            # 构建用户输入
            user_input = {
                "document_content": doc_content,
                "request": "请根据这些文档内容生成一份完整的试卷，包含多种题型，并确保试题与文档内容紧密相关。"
            }

            try:
                # 调用AI接口
                response = client.chat.completions.create(
                    model="qwen2.5-14b-instruct-1m",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": str(user_input)}
                    ],
                    stream=False
                )

                # 获取AI生成的试题内容
                exam_content = response.choices[0].message.content

                # 创建Word文档
                doc = Document()
                doc.add_heading(f"{project.name} - 试题", 0)

                # 添加试题内容
                for paragraph in exam_content.split('\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)

                # 保存文档
                static_dir = os.path.join(settings.BASE_DIR, 'static', 'uploads')
                if not os.path.exists(static_dir):
                    os.makedirs(static_dir)

                file_name = f"{project.id}_{project.name}_试题.docx"
                file_path = os.path.join(static_dir, file_name)

                # 检查文件是否已存在，如果存在则添加时间戳
                if os.path.exists(file_path):
                    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                    file_name = f"{project.name}_试题_{timestamp}.docx"
                    file_path = os.path.join(static_dir, file_name)

                doc.save(file_path)

                # 获取文件URL
                file_url = os.path.join('static', 'uploads', file_name)

                # 保存文件信息到数据库
                file = File.objects.create(
                    name=file_name,
                    file_url="http://127.0.0.1:8000/" + file_url,
                    file_path=file_path,
                    user=user
                )

                # 保存文件与项目的关联
                FileProjectAssociation.objects.create(
                    file=file,
                    project=project
                )

                return Suc()

            except Exception as e:
                print(f"调用AI或保存文件时出错: {str(e)}")
                return Fail(f"生成试题失败: {str(e)}")

        except json.JSONDecodeError:
            return Fail("无效的JSON数据")
        except Exception as e:
            return Fail(f"处理请求时出错: {str(e)}")

    return Fail("无效的请求方法")